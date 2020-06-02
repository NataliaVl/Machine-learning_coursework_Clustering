import lxml as lxml
import pymongo
import hashlib
import os
import connection
from bs4 import BeautifulSoup
import docx

def parsing_doc(path):
    f = open(path, 'r', encoding="utf-8")
    title = f.readline()

    doc = title + f.read()
    _line = doc.partition('Содержание статьи:\n')
    _l = _line[2].partition('\n\n')
    description = _l[0]
    f.close()

    import_doc(doc, title, description)


def parsing_html(path):
    with open(path, "r", encoding="utf-8") as f:
        contents = f.read()
        soup = BeautifulSoup(contents, 'lxml')
        title = soup.title.text + '\n'
        description = soup.aside.text
        doc = title + '\n' + soup.main.text

        import_doc(doc, title, description)


def parsing_docx(path, title):
    os.chdir(path)
    doc1 = docx.Document(title)
    title = doc1.paragraphs[0].text
    description = doc1.paragraphs[4].text
    text = []
    for paragraph in doc1.paragraphs:
        text.append(paragraph.text)
    doc = '\n'.join(text)

    import_doc(doc, title, description)


def import_doc(doc, title, description):
    hash_object = hashlib.md5(doc.encode())
    hash = hash_object.hexdigest()
    try:
        mydict = {"_id": hash, "doc": doc, "title": title, "description": description}
        x = col.insert_one(mydict)
    except Exception as e:
        print(e)



col = connection.connection_mongodb_texts()

directory = 'C://archive_storage_for_clustering'
files = os.listdir(directory)

for item in files:
    name = item.partition('.')
    doc_format = name[2]
    way = directory + "/" + item
    if doc_format == 'html':
        parsing_html(way)
    if doc_format == 'txt':
        parsing_doc(way)
    if doc_format == 'docx':
        parsing_docx(directory, item)